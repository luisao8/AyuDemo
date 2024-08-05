import streamlit as st
from openai import OpenAI
import time
import logging
import PyPDF2
from anthropic import Anthropic
import json
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import requests
import firebase_admin
from firebase_admin import initialize_app
from firebase_admin import credentials
from firebase_admin import storage
from io import BytesIO
import datetime

anthropic = Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])
open_ai = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

firebase_creds = dict(st.secrets["firebase"])
with open("firebase_credentials.json", "w") as f:
    json.dump(firebase_creds, f)

cred = credentials.Certificate("firebase_credentials.json")
app = firebase_admin.initialize_app(cred, {'storageBucket':  'ayudasdemo.appspot.com'})
bucket = storage.bucket()
# Return the path, not the content

model_anthropic = "claude-3-5-sonnet-20240620"

preguntas_memoria_proyecto = [
    {
        "id": 1,
        "titulo": "Actividad que desarrolla la empresa",
        "descripcion": (
            "Ubicación del centro donde se realicen las inversiones. "
            "Descripción somera de las instalaciones actuales del centro productivo objeto del proyecto y "
            "efectos positivos esperados para el desarrollo de la zona (si procede)."
        ),
        "ejemplo": (
            """Adapta es una empresa que tiene como misión ofrecer un servicio de consultoría profesional, donde la elevada
                calidad de nuestros productos y servicios sea la base de nuestra diferenciación, con el claro objetivo de ayudar
                a las empresas a hacer de la innovación un factor clave en la mejora de su competitividad. Hallar soluciones
                globales a las necesidades en términos de innovación de nuestros clientes es nuestra prioridad.
                Nuestra visión es transformar el conocimiento sobre la tecnología, digitalización e innovación en VALOR para
                nuestros clientes mediante servicios de calidad y excelencia en tecnología, permitiendo un crecimiento
                sostenible de la empresa y el desarrollo de las competencias tecnológicas de nuestros trabajadores.
                Nuestro centro de trabajo se ubica en la Calle Arcedianos 20 de Siguenza (Guadalajara). El objeto de nuestros
                servicios se centra en la gestión integral y proactiva de todas las fuentes de financiación para la investigación
                desarrollo e innovación, también en la definición de estrategias avanzadas de innovación y en la
                sistematización de los procesos para su fomento dentro de las empresas y la sociedad. Este asesoramiento
                se personaliza según la tipología, dimensión y necesidad de la empresa. Atendiendo a ello uno de nuestras
                actividades será mentorizar en estas actividades a las empresas de la zona de Guadalajara"""
        )
    },
    {
        "id": 2,
        "titulo": "Descripción del proyecto y de las inversiones a realizar",
        "descripcion": "Explicar los pasos y estrategias que se seguirán para ejecutar el proyecto.",
        "ejemplo": (
            """Nuestro proyecto se enfoca en ayudar a otras empresas a adaptarse y aprovechar al máximo las
                oportunidades que ofrecen las tecnologías digitales y las innovaciones. Para lograrlo, seguimos los siguientes
                pasos:
                Investigación y análisis: Realizamos una investigación detallada de las tendencias y tecnologías más actuales
                del mercado, con el objetivo de identificar aquellas que tienen mayor potencial de transformar los modelos de
                negocio y las prácticas de la industria.
                Análisis de las necesidades de los clientes: Hablamos con los clientes para conocer sus objetivos, metas y
                desafíos empresariales. Este análisis nos permite comprender las necesidades de la empresa y determinar
                la mejor estrategia para ayudarles a alcanzar sus objetivos.
                Diseño de soluciones personalizadas: Con base en los resultados de la investigación y el análisis de
                necesidades de los clientes, desarrollamos soluciones personalizadas para cada cliente. Esto incluiría
                estrategias de digitalización y automatización, así como el desarrollo de nuevos modelos de negocio y la
                optimización de los existentes.
                Implementación y seguimiento: Ayudamos a implementar las soluciones diseñadas, trabajando en estrecha
                colaboración con los equipos internos de la empresa del cliente. También llevaría a cabo un seguimiento
                constante para asegurarse de que los cambios implementados están teniendo el impacto deseado y ajustarlos
                en caso de ser necesario.
                Capacitación y soporte continuo: Ofrecemos capacitación y soporte continuo a los equipos internos del cliente,
                para que puedan adoptar nuevas tecnologías y procesos de manera efectiva y asegurarse de que los cambios
                se mantengan en el tiempo.

                ANEXO III: MEMORIA DEL PROYECTO EMPRESARIAL A REALIZAR
                LINEA DE SUBVENCIÓN PARA EL FOMENTO DE LA INVERSIÓN Y LA MEJORA DE LA
                PRODUCTIVIDAD EMPRESARIAL EN CASTILLA-LA MANCHA (FIE-2023)

                2
                En conclusión, ayudamos a las empresas a adaptarse a los cambios en el mercado, a aprovechar las
                oportunidades que ofrecen las nuevas tecnologías y a mejorar su eficiencia y competitividad a través de la
                digitalización.
                La principal inversión que vamos a realizar en los inicios es la adquisición del material Hardware y Software
                para los dos socios de la empresa. El material HW está compuesto de 2 portátiles, 1 móvil y 2 tablets y el SW
                esta compuesta por 2 licencias de Office y una licencia de un ERP especifico para la empresa."""
        )
    },
    {
        "id": 3,
        "titulo": "Descripción del proceso de producción y del producto final",
        "descripcion": (
            "Necesidades que cubre y mejoras previstas. "
            "Incremento previsto de la facturación. "
            "En el caso de los proyectos de ampliación de la capacidad productiva, indicar el número de "
            "unidades producidas o de prestación de servicios por unidad de tiempo (hora, día, mes o bien anual), "
            "antes y después de llevarse a cabo el proyecto de inversión."
        ),
        "ejemplo": (
            """El material que necesitamos nos va a permitir integrar y automatizar los procesos empresariales clave en un
            único sistema centralizado. Los beneficios incluyen:
            • Mejora de la eficiencia: Al integrar todos los procesos empresariales en un único sistema, los datos
            se actualizan automáticamente y se comparten en toda la organización, lo que reduce la duplicación
            de datos y la necesidad de ingreso manual de información. Esto, a su vez, reduce el tiempo y los
            recursos necesarios para realizar tareas administrativas.
            • Mayor control y visibilidad: Nos proporciona una visión completa y detallada de todas las
            operaciones empresariales en tiempo real. Esto permite a los gerentes supervisar el progreso de los
            procesos en tiempo real, detectar problemas y tomar decisiones informadas y oportunas.
            • Reducción de costos: Nos reduce los costos de una empresa al optimizar los procesos, eliminar
            tareas innecesarias y minimizar la necesidad de contratar personal adicional para administrar los
            procesos.
            • Mejora de la toma de decisiones: Nos proporciona informes precisos y actualizados en tiempo real,
            lo que permite a los gerentes tomar decisiones informadas y oportunas basadas en datos en lugar
            de tener que confiar en conjeturas o suposiciones.
            • Mayor flexibilidad y escalabilidad: Nos satisface las necesidades específicas y se adaptar a medida
            que cambian nuestras necesidades.
            Con esta infraestructura se tiene previsto aumentar la gestión de los expedientes en un 30 % con respecto a
            trabajar sin estas herramientas."""
        )
    },
    {
        "id": 4,
        "titulo": "Creación o mantenimiento de empleo que implica la puesta en marcha del proyecto",
        "descripcion": (
            "Indicar los puestos de trabajo a crear en el centro de realización de las inversiones "
            "(deben coincidir con los indicados en el Anexo I-Solicitud de ayuda) así como de las jornadas que "
            "tendrán cada uno de ellos."
        ),
        "ejemplo": (
            """En principio actualmente no esta trabajando ningún empleado por cuenta ajena, pero a medida que se
desarrolle el ejercicio 2023 y si las previsiones de crecimiento se cumplen durante el presente ejercicio se
contratara a una persona para cubrir el aumento de la carga de los expedientes que se generarán."""
        )
    },
    {
        "id": 5,
        "titulo": "Implicaciones y mejoras medioambientales que implica el proyecto",
        "descripcion": (
            "Indicar si el proyecto está destinado o incorpora la reducción de vertidos a las aguas superficiales "
            "y subterráneas y/o evita contaminación atmosférica, así como en su caso, indicar si el proyecto "
            "incorpora o favorece la recuperación del entorno o la reducción, recuperación o tratamiento de residuos industriales."
        ),
        "ejemplo": (
            """El impacto ambiental que tiene la actividad de la empresa es nulo y la adquisición de estos activos
tecnológicos van a minimizar mas el desarrollo de su actividad."""
        )
    },
    {
        "id": 6,
        "titulo": "Presupuesto desglosado del coste del proyecto",
        "descripcion": "Proveer un desglose detallado del coste del proyecto.",
        "ejemplo": (
            """Hardware (2 Portatiles, 2 tablets, 1 móvil): 4.524,97 + IVA
2 licencias Office: 440 € + IVA
ERP Gestión: 2.988,00€ + IVA"""
        )
    },
    {
        "id": 7,
        "titulo": "Cuenta de pérdidas y ganancias previsional a 3 años",
        "descripcion": (
            "Proveer una previsión de facturación y resultados antes de impuestos para los próximos 3 años."
        ),
        "ejemplo": (
            """Adapta es una empresa que se ha creado en Enero del 2023 y tiene una previsión de facturación de 100.000
€ en el 2023, teniendo previsto un resultado antes de impuestos de 10.000 €. Para los posteriores años se
tiene previsto un crecimiento del 25 % en el 2024 y de un 35 % en el 2025."""
        )
    },
    {
        "id": 8,
        "titulo": "Fuentes de financiación del proyecto con especial referencia a financiación ajena",
        "descripcion": (
            "Detallar las fuentes de financiación del proyecto, destacando especialmente la financiación ajena."
        ),
        "ejemplo": (
            """Este proyecto esta financiado con fondos propios de los accionistas, no habiendo otras fuentes de financiación
externas."""
        )
    },
    {
        "id": 9,
        "titulo": "Análisis comercial del mercado al que va dirigido",
        "descripcion": (
            "Con las previsiones de clientela, contratos o compromisos suscritos con posibles clientes. "
            "Características del mercado y del cliente. Situación del mercado al que se destinan los productos, "
            "análisis de la competencia, ventajas comparativas, expectativas y futuras vías de comercialización. "
            "En el caso de ampliaciones o transformación fundamental del proceso global de producción se deben "
            "suministrar datos antes de la inversión y previsiones una vez realizada la inversión sobre unidades "
            "vendidas y valor de éstas."
        ),
        "ejemplo": (
            """El mercado de la consultoría de innovación y digitalización está experimentando un crecimiento significativo
a nivel mundial debido a la creciente demanda de empresas que buscan aprovechar los beneficios que
ofrecen las nuevas tecnologías y la necesidad de adaptarse a un entorno empresarial en constante cambio.
A continuación presentamos algunos aspectos clave del análisis comercial de este mercado:
Crecimiento del mercado: Se espera que siga creciendo en el futuro debido a la creciente demanda de las
empresas para mejorar la eficiencia operativa, desarrollar nuevos modelos de negocio y aprovechar las
oportunidades que ofrecen las nuevas tecnologías. Según un informe de MarketsandMarkets, se espera que
el mercado global de la consultoría de gestión alcance los 200 mil millones de dólares en 2025, con una tasa
de crecimiento anual compuesta (CAGR) del 6,2% durante el período 2020-2025.
Competencia: El mercado es altamente competitivo, con un gran número de empresas de consultoría que
ofrecen servicios similares. Las empresas líderes del mercado incluyen McKinsey, Accenture, Deloitte, PwC,
Bain, entre otros.
Tendencias del mercado: Algunas de las tendencias clave en el mercado de la consultoría incluyen la
adopción de soluciones basadas en la nube, la inteligencia artificial (IA) y el análisis de datos, el enfoque en
la transformación digital, la incorporación de la sostenibilidad y la preocupación por la ciberseguridad.
Demanda regional: La demanda de estos servicios varía según las regiones del mundo. Actualmente, América
del Norte es el mercado más grande para la consultoría, seguido de Europa y Asia Pacífico.
Clientes: Los clientes de la consultoría incluyen empresas de diversos sectores, desde la industria
manufacturera hasta el sector financiero y de servicios. Estos clientes buscan mejorar la eficiencia, desarrollar
nuevos modelos de negocio y aprovechar las oportunidades que ofrecen las nuevas tecnologías para competir
en el mercado.

En conclusión, el mercado está experimentando un fuerte crecimiento debido a la creciente demanda de
servicios que ayuden a las empresas a adaptarse a un entorno empresarial en constante cambio y aprovechar
las oportunidades que ofrecen las nuevas tecnologías. Sin embargo, la competencia es alta y las tendencias
del mercado están evolucionando rápidamente, lo que requiere que las empresas como Adapta estén
constantemente innovando y actualizándose para seguir siendo relevantes.
En relación a la previsión de clientes para este ejercicio 2023, se tiene previsto superar los 20 clientes con un
ticket medio de 5000 €. Nuestra tipología de cliente es pequeña empresa con facturación entre 1 y 5 millones
de €."""
        )
    },
    {
        "id": 10,
        "titulo": "Calendario previsto de ejecución del proyecto",
        "descripcion": "Desarrollar lo indicado en el Anexo I.",
        "ejemplo": (
            """La ejecución del proyecto se va a desarrollar entre Febrero y Mayo del 2023."""
        )
    },
    {
        "id": 11,
        "titulo": "Información sobre la necesidad de la ayuda y su efecto esperado",
        "descripcion": (
            "Breve explicación de la necesidad de la ayuda y de su efecto en la decisión de invertir o en la decisión "
            "en cuanto al emplazamiento. Debe incluirse una explicación de la inversión o del emplazamiento alternativos "
            "en caso de que no se conceda la ayuda, dependiendo del supuesto que se encuentre entre: "
            "1) la ayuda ofrece un incentivo para tomar la decisión de invertir en la zona en cuestión porque, de otro "
            "modo, la inversión no sería lo suficientemente rentable para el beneficiario de la ayuda en ningún lugar del "
            "Espacio Económico Europeo (supuesto 1, decisión de invertir), "
            "2) la ayuda ofrece un incentivo para localizar la inversión prevista en la zona en cuestión, en vez de "
            "hacerlo en otro lugar, porque compensa las desventajas y los costes netos de invertir en un emplazamiento "
            "situado en dicha zona (supuesto 2, decisión en cuanto al emplazamiento)."
        ),
        "ejemplo": (
            """La ayuda es un incentivo adicional para monitorizar todos sus servicios y tomar la decisión de invertir en otras
actividades que nos permitan cumplir con el crecimiento previsto en este ejercicio."""
        )
    }
]

system_prompt = f""""
Contexto: 

Estás actuando como un asistente inteligente especializado en completar la Memoria del Proyecto Empresarial. Tu tarea es generar respuestas detalladas, profesionales y que aporten valor a cada pregunta del formulario, utilizando los ejemplos proporcionados como referencia. Debes mejorar y expandir estas respuestas, asegurándote de que sean coherentes, persuasivas y bien estructuradas.

Información Disponible:

Texto extraído del PDF de Información de la Empresa: Este texto proporciona un contexto general sobre la empresa, sus actividades, misión, visión, y cualquier otra información relevante para entender su operación y objetivos.

Datos Financieros en Formato JSON de las Cuentas Anuales: Estos datos incluyen indicadores financieros clave como el volumen de negocio, número de empleados, resultados del ejercicio, activos totales, y cualquier otro dato financiero relevante.

Problem Statement Generado por IA: Este statement ofrece un resumen de los desafíos y oportunidades específicos que el proyecto empresarial busca abordar, basado en una conversación previa con el usuario.

Objetivos de la IA:

Mejora de Calidad: Superar la calidad de los ejemplos proporcionados, generando respuestas más completas, persuasivas y bien articuladas.

Integración de Datos: Utilizar los datos financieros y el contexto del problem statement para enriquecer las respuestas y proporcionar justificaciones cuantitativas y cualitativas.

Claridad y Estructura: Asegurar que cada respuesta sea clara, fácil de leer y estructurada de manera lógica para facilitar la comprensión del lector.

Enfoque en el Valor: Destacar el valor añadido que el proyecto y la empresa aportan a la comunidad y al mercado objetivo, enfatizando la innovación, sostenibilidad y competitividad.

Proceso para Responder:

Leer la Pregunta y Descripción: Comienza leyendo detenidamente la pregunta y su descripción para comprender exactamente lo que se está solicitando.

Analizar el Ejemplo Proporcionado: Revisa el ejemplo proporcionado y utiliza como guía para identificar áreas de mejora y expansión en la respuesta.

Integrar Información Relevante: Incorpora información relevante de los datos financieros y del problem statement para respaldar tus respuestas con evidencias y análisis detallados.

Formular Respuestas Mejoradas:

Estructura la respuesta en párrafos claros y lógicos.
Utiliza lenguaje profesional y persuasivo.
Resalta el valor añadido del proyecto para la empresa y la comunidad.
Asegúrate de que la respuesta sea comprensible y atractiva para el lector.
Revisión y Refinamiento: Antes de finalizar, revisa la respuesta para verificar su coherencia, precisión y claridad. Haz ajustes necesarios para mejorar la presentación y el impacto de la respuesta.
"""

def generate_timestamp():
    return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

def generate_temporal_url(document_buffer, document_type):
    timestamp = generate_timestamp()
    blob = bucket.blob(f"{document_type}_{timestamp}.docx")
    blob.upload_from_file(document_buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    url_temporal = blob.generate_signed_url(version='v4', expiration=3600, method='GET')
    return url_temporal

def format_json_to_html(memoria_url, declaracion_url):
    html_content = "<h3>URLs de los informes:</h3>"
    html_content += "<h4>Enlaces de descarga:</h4><ul>"
    html_content += f"<li><a href='{memoria_url}' style='color: blue;'>Descargar Memoria del Proyecto</a></li>"
    html_content += f"<li><a href='{declaracion_url}' style='color: blue;'>Descargar Declaración Responsable</a></li>"
    html_content += "</ul>"
    html_content += "<p><br>Atentamente, <br>Asistente IA de Generación de Informes.</p>"
    return html_content

def send_mail_to_make(email, subject, html_email):
    request_body = {
        "email": email,
        "subject": subject,
        "body": html_email
    }
    response = requests.post("https://hook.eu2.make.com/9hid97sxq5gfw72ilcftv0w9yh45ioby", json=request_body)
    return response

def extract_text_from_pdf(uploaded_file):
    if uploaded_file is None:
        logging.error("Error: No file uploaded")
        return None
    
    try:
        pdf_file = io.BytesIO(uploaded_file.getvalue())
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_content = ""
        for page in pdf_reader.pages:
            text_content += page.extract_text() + "\n"
        return text_content
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {e}")
        return None
    
def extract_financial_data(extracted_text):
    try:
        prompt = f"""
        Eres un asistente especializado en extraer información financiera de documentos españoles.
        Analiza el siguiente documento de Cuentas Anuales de una empresa española y extrae la siguiente información en formato JSON:

        {{
            "empresa": {{
                "nif": "",
                "razon_social": "",
                "total_trabajadores": "",
                "volumen_negocio": "",
                "balance_general_anual": ""
            }},
            "año": "",
            "presentante": {{
                "nombre_apellidos": "",
                "dni_nif_pasaporte": ""
            }}
        }}

        Asegúrate de rellenar todos los campos con la información correspondiente del documento. Si algún dato no está disponible, déjalo en blanco.

        Documento de Cuentas Anuales:
        {extracted_text}
        """

        response = anthropic.messages.create(
            model=model_anthropic,
            max_tokens=1000,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )

        content = response.content[0].text if isinstance(response.content, list) else response.content
        logging.debug("Raw content: %s", content)  # For debugging

        json_start = content.find('{')
        json_end = content.rfind('}') + 1
        if json_start != -1 and json_end != -1:
            json_str = content[json_start:json_end]
            return json.loads(json_str)
        else:
            raise ValueError("No se pudo extraer un JSON válido de la respuesta")

    except Exception as e:
        logging.error(f"Error extracting financial data: {e}")
        return None

def generate_prompt(financial_data, company_description, goal_statement, system_prompt, pregunta, previous_responses):
    context = "\n".join([f"Pregunta {k}: {v}" for k, v in previous_responses.items()])

    formatted_prompt = f"""
# Contexto:

Estás actuando como un asistente inteligente especializado en completar la Memoria del Proyecto Empresarial para una solicitud de ayuda en Castilla-La Mancha. El objetivo de esta ayuda es:

Facilitar financiación a las empresas para poner en marcha un proyecto empresarial, su incorporación al tejido empresarial de Castilla-La Mancha y la mejora de su productividad, todo ello mediante la realización de inversiones productivas, la mejora de los procesos y productos fabricados en la región, la diferenciación de estos frente a la competencia y la apertura a nuevos mercados para alcanzar la consolidación del tejido empresarial de la región.

Tu tarea es generar respuestas detalladas, profesionales y que aporten valor a cada pregunta del formulario, utilizando los ejemplos proporcionados como referencia. Debes mejorar y expandir estas respuestas, asegurándote de que sean coherentes, persuasivas y bien estructuradas.

# Información Disponible:

Texto extraído del PDF de Información de la Empresa: 
{company_description}

Datos Financieros en Formato JSON de las Cuentas Anuales:
{json.dumps(financial_data, indent=2, ensure_ascii=False)}

Problem Statement Generado por IA:
{goal_statement}

# Objetivos de la IA:

- Mejora de Calidad: Superar la calidad de los ejemplos proporcionados, generando respuestas más completas, persuasivas y bien articuladas.
- Integración de Datos: Utilizar los datos financieros y el contexto del problem statement para enriquecer las respuestas y proporcionar justificaciones cuantitativas y cualitativas.
- Claridad y Estructura: Asegurar que cada respuesta sea clara, fácil de leer y estructurada de manera lógica para facilitar la comprensión del lector.
- Enfoque en el Valor: Destacar el valor añadido que el proyecto y la empresa aportan a la comunidad y al mercado objetivo, enfatizando la innovación, sostenibilidad y competitividad.
- Coherencia Global: Asegurarse de que todas las respuestas sean coherentes entre sí y formen un conjunto cohesivo que respalde la solicitud de ayuda.

# Proceso para Responder:
- Leer la Pregunta y Descripción: Comienza leyendo detenidamente la pregunta y su descripción para comprender exactamente lo que se está solicitando.
- Revisar el Contexto Previo: Lee las respuestas anteriores para asegurar la coherencia y evitar repeticiones.
- Analizar el Ejemplo Proporcionado: Utiliza el ejemplo como guía para identificar áreas de mejora y expansión en la respuesta.
- Integrar Información Relevante: Incorpora información relevante de los datos financieros y del problem statement para respaldar tus respuestas con evidencias y análisis detallados.
- Formular Respuestas Mejoradas:
  - Estructura la respuesta en párrafos claros y lógicos.
  - Utiliza lenguaje profesional y persuasivo.
  - Resalta el valor añadido del proyecto para la empresa y la comunidad de Castilla-La Mancha.
- Revisión y Refinamiento: Antes de finalizar, revisa la respuesta para verificar su coherencia, precisión y claridad. Haz ajustes necesarios para mejorar la presentación y el impacto de la respuesta.
- No te inventes datos financieros, utiliza los datos del json.

## Contexto de respuestas previas:
{context}

{system_prompt}

Pregunta actual:
{pregunta['titulo']}
{pregunta['descripcion']}

Por favor, proporciona una respuesta detallada y profesional para esta pregunta, mejorando el ejemplo dado y utilizando la información proporcionada sobre la empresa y sus datos financieros. Asegúrate de que tu respuesta sea coherente con las respuestas anteriores y contribuya a una narrativa global convincente para la solicitud de ayuda. 
Solo devuelva la respuesta directa, sin comentarios. Por ejemplo... "Grupo Aire Limpio es...". Así no habrá que formatear más allá de tus respuestas.
"""
    return formatted_prompt

def generate_response(prompt):
    anthropic = Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])
    response = anthropic.messages.create(
        model=model_anthropic,
        max_tokens=1000,
        messages=[
            {"role": "user", "content": prompt}
        ]
    )
    return response.content[0].text if isinstance(response.content, list) else response.content

def extract_razon_social(flattened_json):
    """
    Extract the 'razon_social' from the flattened JSON and remove it from the dictionary.
    
    :param flattened_json: The flattened JSON object
    :return: A tuple containing the extracted 'razon_social' and the modified JSON
    """
    razon_social = flattened_json.get('empresa_razon_social', '')
    
    # Remove the 'razon_social' key from the flattened JSON
    flattened_json.pop('empresa_razon_social', None)
    
    return razon_social, flattened_json

def generate_problem_statement(thread_id):
    
    assistant_id = st.secrets["PROBLEM_STATEMENT_AI"]  # Replace with your actual assistant ID

    # Create a new message in the thread asking for a problem statement
    message = open_ai.beta.threads.messages.create(
        thread_id=thread_id,
        role="user",
        content="Basandonos en nuestra conversación, genera un goal_statement para el proyecto. Conciso y directo:"
    )

    # Run the assistant
    run = open_ai.beta.threads.runs.create(
        thread_id=thread_id,
        assistant_id=assistant_id
    )

    # Wait for the run to complete
    while run.status != "completed":
        time.sleep(1)
        run = open_ai.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)

    # Retrieve the assistant's response
    messages = open_ai.beta.threads.messages.list(thread_id=thread_id)
    problem_statement = messages.data[0].content[0].text.value

    return problem_statement

goal_statement_text = """Problem Statement:

Grupo Aire Limpio busca financiación para mejorar su infraestructura tecnológica mediante la adquisición de nuevos servidores internos. Esta inversión permitirá a la empresa optimizar el procesamiento de datos de las Unidades de Tratamiento de Aire (UTAs) que gestiona, aumentando la eficiencia operativa y reduciendo significativamente el consumo energético.

Propósito del Proyecto:

La subvención solicitada se destinará a la compra de servidores de última generación que facilitarán el manejo de grandes volúmenes de información, mejorando así la capacidad de procesamiento y análisis de datos críticos relacionados con la calidad del aire interior en edificios comerciales y hospitales. Este proyecto es fundamental para mantener la posición de liderazgo de Grupo Aire Limpio en el sector de la sostenibilidad y el bienestar ambiental, al garantizar que sus soluciones de aire interior sean más eficientes y sostenibles.

Plan de Acción:

El plan de acción incluye la adquisición de servidores de alta eficiencia energética y su implementación en los centros de datos de la empresa. Se prevé que estos nuevos servidores mejoren la capacidad de procesamiento de información de las UTAs en un 40%, lo que permitirá un monitoreo más preciso y en tiempo real de la calidad del aire interior. Además, el plan contempla la capacitación del personal técnico para asegurar el uso óptimo de las nuevas tecnologías implementadas.

Costes del Proyecto:

El coste total del proyecto se estima en 40,000 euros, que cubrirán la compra de equipos, la instalación de los servidores y la capacitación del personal. Esta inversión se alinea con los objetivos de sostenibilidad del Grupo Aire Limpio, al reducir el consumo energético en un 20% y minimizar la huella de carbono de sus operaciones.

Proyección de Crecimiento e Impacto:

La implementación de este proyecto no solo mejorará la eficiencia operativa de Grupo Aire Limpio, sino que también se espera que impulse un crecimiento del 15% en la facturación anual, gracias a la mejora en la calidad del servicio y la capacidad de atender a un mayor número de clientes. Además, se proyecta que el proyecto contribuirá al desarrollo económico de Castilla-La Mancha al generar empleos y fomentar la adopción de prácticas sostenibles en la región.

Contribución a la Sociedad Manchega:

Este proyecto refuerza el compromiso de Grupo Aire Limpio con la sostenibilidad y el bienestar de la comunidad local. Al mejorar la eficiencia energética de sus operaciones, la empresa reducirá el impacto ambiental y promoverá un entorno más saludable para los habitantes de Castilla-La Mancha. Además, la iniciativa impulsará la creación de empleos locales, ya que se espera contratar a técnicos y especialistas para gestionar los nuevos sistemas implementados."""

def transform_financial_data_keys(financial_data):
    key_mapping = {
        'empresa_nif': '[NIF]',
        'empresa_razon_social': '[razon_social]',
        'presentante_dni_nif_pasaporte': '[id_rep_legal]',
        'presentante_nombre_apellidos': '[nombre_rep_legal]',
        'año': '[ano_cierre]',
        'empresa_total_trabajadores': '[num_trabajadores]',
        'empresa_volumen_negocio': '[vol_negocio]',
        'empresa_balance_general_anual': '[balance_general]'
    }
    
    transformed_data = {}
    for old_key, value in financial_data.items():
        new_key = key_mapping.get(old_key, old_key)
        transformed_data[new_key] = value
    
    return transformed_data

def prepare_info_json(responses):
    """
    Prepare the information JSON with wrapped keys and single-level values.
    """
    return {f"[{key}]": value for key, value in responses.items()}

def flatten_json(nested_json, prefix=''):
    """
    Flatten a nested JSON object.
    """
    flattened = {}
    for key, value in nested_json.items():
        if isinstance(value, dict):
            flattened.update(flatten_json(value, f"{prefix}{key}_"))
        else:
            flattened[f"{prefix}{key}"] = value
    return flattened

def apply_format(run, font_name='Lucida', font_size=9):
    run.font.name = font_name
    run.font.size = Pt(font_size)

def replace_text_in_paragraph(paragraph, json_data):
    for key, value in json_data.items():
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, str(value))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                apply_format(run)

def process_table(table, json_data):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, json_data)
            # Check for nested tables
            for nested_table in cell.tables:
                process_table(nested_table, json_data)

def doc_from_firebase(blob_name):
    blob = bucket.blob(blob_name)
    blob_bytes = blob.download_as_bytes()
    return blob_bytes

def fill_docx_template(template_blob_name, json_data):
    template_bytes = doc_from_firebase(template_blob_name)
    doc = Document(BytesIO(template_bytes))
    
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, json_data)
    
    for table in doc.tables:
        process_table(table, json_data)
    
    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer

def generar_contrato(email, info_pdf, annual_accounts, thread_id):
    logging.info("Starting contract generation process")
    
    # Extract text from PDFs
    logging.info("Extracting text from PDFs")
    company_description_text = extract_text_from_pdf(info_pdf)
    financial_info_text = extract_text_from_pdf(annual_accounts)

    # Generate problem statement
    logging.info("Generating problem statement")
    goal_statement_text = generate_problem_statement(thread_id)

    # Extract financial data
    logging.info("Extracting financial data")
    financial_data_json = extract_financial_data(financial_info_text)
    financial_data_json_for_word = financial_data_json.copy()
    # Extract razon_social and create two separate JSON objects
    razon_social, financial_data_json_without_razon_social = extract_razon_social(financial_data_json)

    financial_data_json_for_word['empresa']['[razon_social]'] = razon_social

    # Initialize an empty dictionary to store responses
    all_responses = {}

    # Generate prompts and get responses from Claude iteratively
    logging.info("Starting response generation for each question")
    for i, pregunta in enumerate(preguntas_memoria_proyecto, start=1):
        logging.info(f"Generating prompt for question {i}")
        prompt = generate_prompt(
            financial_data_json_without_razon_social, 
            company_description_text, 
            goal_statement_text, 
            system_prompt, 
            pregunta,
            all_responses
        )
        
        # Generate response for the current question
        logging.info(f"Generating response for question {i}")
        response = generate_response(prompt)
        
        # Add the new response to all_responses
        all_responses[f"[num_{i}]"] = response
        logging.info(f"Response generated and stored for question {i}")

    # Prepare information JSON for Memoria del Proyecto
    info_json = all_responses
    logging.info("Prepared info_json for Memoria del Proyecto")

    # Transform and flatten financial data for Declaración Responsable
    logging.info("Transforming and flattening financial data")
    flat_financial_data = flatten_json(financial_data_json_for_word)
    transformed_financial_data = transform_financial_data_keys(flat_financial_data)

    # Fill the first document (Memoria del Proyecto)
    logging.info("Filling Memoria del Proyecto document")
    template_blob_memoria = "placeholders_memoria.docx"
    memoria_buffer = fill_docx_template(template_blob_memoria, info_json)

    # Fill the second document (Declaración Responsable)
    logging.info("Filling Declaración Responsable document")
    template_blob_declaracion = "placeholders_financiero.docx"
    declaracion_buffer = fill_docx_template(template_blob_declaracion, transformed_financial_data)

    # Generate temporal URLs for both documents
    logging.info("Generating temporal URLs for documents")
    memoria_url = generate_temporal_url(memoria_buffer, "memoria_proyecto")
    declaracion_url = generate_temporal_url(declaracion_buffer, "declaracion_responsable")

    # Format HTML email content
    logging.info("Formatting HTML email content")
    html_email = format_json_to_html(memoria_url, declaracion_url)

    # Send email via make.com
    logging.info("Sending email via make.com")
    subject = "Documentos generados para su solicitud de ayuda"
    response = send_mail_to_make(email, subject, html_email)

    if response.status_code == 200:
        logging.info("Email sent successfully")
        return "Documents generated and email sent successfully"
    else:
        logging.error(f"Failed to send email. Status code: {response.status_code}")
        return "Documents generated successfully, but failed to send email"








